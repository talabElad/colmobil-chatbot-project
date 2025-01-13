import pandas as pd
import json
import numpy as np
import os
# importing necessary functions from dotenv library
from dotenv import load_dotenv, dotenv_values
import numpy as np
from langchain.agents import AgentExecutor
from pydantic import __init__
from sqlalchemy import create_engine, inspect
import json
from langchain.agents.tool_calling_agent.base import create_tool_calling_agent
import threading
import numpy as np
from langchain.sql_database import SQLDatabase
from langchain.agents.agent_toolkits.sql.toolkit import SQLDatabaseToolkit
from langchain.prompts.chat import ChatPromptTemplate
from sqlalchemy import create_engine, MetaData, Table
from langchain_aws import ChatBedrock
from dotenv import load_dotenv
from langchain_core.tools import tool
from tools import search_additional_descriptions, search_models
from docx import Document
import redis
from langchain_aws import ChatBedrockConverse
from langchain_openai import AzureChatOpenAI



load_dotenv(".env")


# endpoint = os.getenv('RDS_ENDPOINT')
# username = os.getenv('RDS_USERNAME')
# password = os.getenv('RDS_PASSWORD')
# database = os.getenv('RDS_DATABASE')

endpoint="database-colmobil.c9owiq2sebpi.us-east-1.rds.amazonaws.com"
username="admin"
password="Bb123456!"
database="databasecolmobil"

# llm = ChatBedrockConverse(
#     model="us.anthropic.claude-3-5-sonnet-20241022-v2:0",
#     temperature=0,
#     region_name="us-east-1",
#     provider="anthropic",
#     stop_sequences = ["|@|@|"]
# )

# llm = ChatBedrockConverse(
#     model="amazon.nova-pro-v1:0",
#     temperature=0,
#     region_name="us-east-1",
#     provider="amazon",
#     stop_sequences = ["|@|@|"]
# )


os.environ["OPENAI_API_TYPE"]="azure"
os.environ["OPENAI_API_VERSION"]="2024-02-15-preview"
os.environ["AZURE_OPENAI_ENDPOINT"]="https://openaiimagetext.openai.azure.com/" # Your Azure OpenAI resource endpoint
os.environ["OPENAI_API_KEY"]="212f3a6ba66d409c8219de169aefec1a" # Your Azure OpenAI resource key
os.environ["AZURE_OPENAI_GPT4O_MODEL_NAME"]="gpt-4o"
os.environ["AZURE_OPENAI_GPT4O_DEPLOYMENT_NAME"]="sahargpt4o"


llm = AzureChatOpenAI(
    api_version="2024-05-01-preview",
    azure_deployment=os.getenv("AZURE_OPENAI_GPT4O_DEPLOYMENT_NAME"),
    temperature=0,stop_sequences=["|@|@|"]
    )

# Create a SQLDatabase object
connection_string = f"mysql+pymysql://{username}:{password}@{endpoint}/{database}"

# Create the SQLAlchemy engine
engine = create_engine(connection_string)

# Use SQLAlchemy MetaData to reflect the database schema
inspector = inspect(engine)

# Get column names for the 'cars_collection' table
columns = inspector.get_columns('cars_collection')

clean_columns_dynamic_fields = []
column_details = []
column_details_lines = []
for column in columns:
    name = column['name']
    if "embeddings" in name:
        continue
    if not name in ["model","car_id","reason","car_web_link","image_url","brand", "basic_price_nis", "additional_description"]:
        clean_columns_dynamic_fields.append(name)
    data_type = column['type']
    enum_values = "N/A"
    if str(data_type) == 'ENUM':
        enum_values = data_type.enums  # Enum values if applicable (returns None if not an enum)
        column_details.append({
        "name": name,
        "data_type": str(data_type),  # Convert to string for better readability
        "enum_values": enum_values if enum_values else "N/A"  # Handle non-enum columns
        })
        column_details_lines.append(f" - Column: {name}, Type: {data_type}, Enum Values: {enum_values}")
        continue
    
    column_details.append({
        "name": name,
        "data_type": str(data_type),  # Convert to string for better readability
        })  
        
    column_details_lines.append(f" - Column: {name}, Type: {data_type}")

print("hi")


column_names = [column['name'] for column in column_details]

class CustomSQLDatabase(SQLDatabase):
    def get_table_info(self, column_names, table_name: str = None) -> str:
        """
        Override to filter specific tables and columns.
        """
        allowed_table = "cars_collection"
        allowed_columns = column_names

        # Restrict to the specified table
        if table_name and table_name != allowed_table:
            return ""

        # Fetch table schema information
        table_info = super().get_table_info(table_name)
        if not table_name:
            return f"Table: {allowed_table}\n"

        # Filter the column details
        filtered_columns_info = "\n".join(
            [line for line in table_info.split("\n") if any(col in line for col in allowed_columns)]
        )
        return f"Table: {allowed_table}\n{filtered_columns_info}"


db = CustomSQLDatabase(engine=engine)

# Create SQLDatabaseToolkit with the database and the language model
sql_toolkit = SQLDatabaseToolkit(db=db, llm=llm)

chat_leading_questions_doc = Document("leading_questions.docx")

database_column_names_hebrew = Document("database_column_names_hebrew.docx")
print(column_details_lines.__str__())
print(clean_columns_dynamic_fields)
class MasterAgent:
    def __init__(self):
        self.r = redis.Redis(host='localhost', port=6379, db=0)
        self.agent_executor = self.create_master_agent()
        self.response = ""
    
    def create_master_agent(self):
        manager_prompt = ChatPromptTemplate.from_messages(
        [
            ("system", f"""
             אתה סוכן חכם למכירת רכבים, שתפקידו לעזור ללקוחות למצוא את הרכב החדש שהכי מתאים להם.
            אתה מציע רכבים אך ורק מתוך הדאטאבייס הפנימי של כלמוביל, ולכן לעולם אינך ממציא דגמים שאינם קיימים בדאטאבייס ה-SQL.
            עליך להיות מקצועי, נעים הליכות ולדבר בעברית תקינה וזורמת. חשוב גם להימנע ממשפטים ארוכים מדי ולהיות תמציתי אך ברור.
            אתה מנומס, תקשורתי, ובמידה המתאימה גם יודע להיות מצחיק, כדי לייצר חוויית שירות חיובית.

            הצעת רכבים: תמיד תשתדל להציע 3 רכבים סופיים מתוך הדאטאבייס, אלא אם בקשות הלקוח מגבילות את החיפוש למספר קטן יותר של רכבים, וזה בסדר אם לא תמצא רכבים שעונים לדרישות.
            הבנת הצרכים: במידת הצורך, עליך "להגדיל ראש" ולהתחשב בנסיבות מיוחדות שהלקוח מציין. לדוגמה, אם ללקוח יש 4 ילדים, נסה להתחשב במספר המושבים ברכב שתציע.
            במקרים שבהם חסר לך מידע על מנת לסנן את הרכבים ל-3 הצעות מתאימות, השתמש בשאלות מנחות כמו:
            {chat_leading_questions_doc.paragraphs}

            הנחיות חשובות לבניית שאילתות SQL:
            תמיד בדוק את שמות העמודות הקיימות בטבלה לפני ביצוע שאילתה.
            הגב את כמות התוצאות המוחזרות למינימום הנדרש.
            לעולם אל תשתמש באפשרות '*' בשאילתה – בחר תמיד שמות עמודות מדויקים.
            שמות העמודות לשימוש:
            {column_details_lines.__str__()}

            פורמט להצעת רכבים:
            כאשר מצאת רכבים מתאימים (בין 1 ל-3), השתמש בפורמט הבא בתגובה:

            לכל הצעה חובה לכלול שני שדות קבועים: car_id ו-reason.
            בשדה "reason" הסבר בשורה אחת מדוע הרכב מתאים ללקוח, תוך התייחסות לצרכים שהציג.
            בנוסף, בחר 3 שדות דינמיים מתוך הרשימה הבאה (מלבד העמודות car_web_link, image_url, brand, model, basic_price_nis, additional_description שאינן רצויות כשדות):
            {clean_columns_dynamic_fields.__str__()}
            אתה לא מוסיף ערך לשדה דינמי, רק את השם של השדה

            מבנה התגובה:

            התחל בהצעות הרכבים עם סימן |||.
            בין פרטי הרכב הפרד עם |.
            כל שדה ורשומת נתון מופרדים בעזרת ,,.
            סיים את התגובה עם סימן סגירה |@|@| ללא טקסט נוסף לאחריו.

            דוגמה:
                    מצאתי רכבים שאני בטוח שיתאימו לך, אתה כמובן יכול להמשיך להכווין אותי:
                    |||
                    car_id:3 ,, reason:סיבה קצרה ++ num_of_doors ,, trunk_capacity_liters ,, safety_system_type |
                    car_id:16 ,, reason:סיבה קצרה ++ num_of_doors ,, trunk_capacity_liters ,, safety_system_type |
                    car_id:7 ,, reason:סיבה קצרה ++ num_of_doors ,, trunk_capacity_liters ,, safety_system_type |@|@|

                    דגשים נוספים:

                    הצע רק רכבים, דגמים ומידע שקיימים בדאטאבייס הפנימי של כלמוביל.
                    אם לא מצאת רכבים מתאימים, הסבר זאת בצורה נעימה ומקצועית, מבלי להמציא הצעות שאינן קיימות.
                    ודא שכל הדגמים, הערכים והנתונים שאתה מציע אכן מאומתים בדאטאבייס.
                    תגוון במשפט הפתיחה של התשובה ובנימוקים בשדה reason.
    
                                """),
            # ("system", "מידע היסטורי רלוונטי של חיפושים קודמים שלך בכדי לחסוך זמן: {context_info}"),
            ("placeholder", "{chat_history}"),
            ("human", "{input}"),
            ("placeholder", "{agent_scratchpad}"),
        ]
        )
        tools_list = sql_toolkit.get_tools()
        tools_list.append(search_models)
        tools_list.append(search_additional_descriptions)
        
        manager_agent = create_tool_calling_agent(
            llm=llm,
            tools=tools_list,
            prompt=manager_prompt
        )
        agent_executor = AgentExecutor(agent=manager_agent, tools=tools_list,return_intermediate_steps=True, verbose=True, max_iterations=15)
        return agent_executor

    # def update_context(self):
    #     response_intermediate_steps = llm.invoke(f"""
    #         אתה סוכן שמתפקידו לסכם את הפעולות הסופיות, החשובות והרלוונטיות של הסוכן בהתבסס ובהקשר לתשובת העוזר החכם. הסיכום שלך צריך להתמקד בעיקר בערכי grocery_type_id, recipe_id,ו-barcode, וכמובן השמות או המבצעים שהם מייצגים.

    #         הנחיות:

    #         סיכום ללא הנמקה: הצג את הסיכום בצורה ממוקדת, ללא הסברים נוספים.
    #         you save informatrion only about items that exists in the 
    #         שימוש במידע מההיסטוריה:

    #         היסטוריה ישנה: old_history_of_actions:{self.context_dict["context_info"]}
    #         היסטוריה חדשה: new:{str(self.response["intermediate_steps"])}
    #         תשובת העוזר החכם:{self.conv[-1]}
    #         הענקת ערכים ל-IDs: ודא שכל ערך ל-IDs (כגון grocery_type_id, recipe_id, barcode) מוצמד נכון ואינו ריק. השתמש אך ורק בערכים הקיימים במאגר, ואל תמציא ערכים חדשים.

    #         דיוק ואמיתות: אם יש ערכים חסרים או מידע לא זמין, ציין את החוסר במקום להמציא ערכים חדשים. השתמש במידע הקיים בלבד ואל תשאיר ערכים ריקים או לא מדויקים.
    #         אל תחזיר מידע ריק לדוגמה:
    #         2. **מלפפון**:
    #             - **grocery_type_id**: לא סופק
    #             - **recipe_id**: לא סופק
    #             - **barcode**: לא סופק context_info
            
    #         """)
    #     print(self.context_dict["context_info"],"context_info")
    #     print(str(self.response["intermediate_steps"]), "intermediate_steps")
    #     print(response_intermediate_steps.content, "update_context")
    #     self.context_dict["chat_history"] = self.conv
    #     self.context_dict["context_info"] = response_intermediate_steps.content
    #     self.r.hset(self.user_id, 'context_dict',json.dumps(self.context_dict)) 
    #     self.r.hset(self.user_id, 'conv', json.dumps(self.conv))
        


    def custom_invoke(self,input:str, user_id:str)-> str:
        os.environ['UserID'] = user_id
        self.user_id = user_id
        self.initialize_conversation()
        self.context_dict["input"]=input
        print(self.context_dict)
        self.response = self.agent_executor.invoke(self.context_dict)
        print(self.response)
        self.conv.append({"role": "user", "content":input})
        self.conv.append({"role": "assistant", "content":self.response['output']})
        self.context_dict["chat_history"] = self.conv
        self.r.hset(self.user_id, 'context_dict',json.dumps(self.context_dict)) 
        self.r.expire(self.user_id+':' + 'context_dict', 172800)
        self.r.hset(self.user_id, 'conv', json.dumps(self.conv)) 
        self.r.expire(self.user_id+':' + 'conv', 172800)
        # threading.Thread(target=self.update_context).start()
        return self.response['output']
    
    def initialize_conversation(self):
        print("1")
        if self.r.hexists(self.user_id, 'conv'):
            self.context_dict = json.loads(self.r.hget(self.user_id, 'context_dict'))  # Set a hash field
            
            self.conv = json.loads(self.r.hget(self.user_id, 'conv'))
        else:
            self.context_dict = {"input":"","chat_history":[],"context_info":""}
            self.conv=[]
        print("2")
        
        

if __name__=="__main__":
    main_agent_executer = MasterAgent()